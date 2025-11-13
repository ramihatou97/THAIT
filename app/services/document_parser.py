"""
NeuroscribeAI - Document Parsing Service
Comprehensive document parsing for PDF, DOCX, and text files with OCR support
"""

import logging
import io
from typing import Optional, Dict, Any, List
from pathlib import Path

# PDF parsing
import fitz  # PyMuPDF

# DOCX parsing
from docx import Document as DocxDocument

# OCR (optional - requires tesseract binary)
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)


# =============================================================================
# Document Parser Service
# =============================================================================

class DocumentParserService:
    """Service for parsing various document formats"""

    def __init__(self):
        """Initialize document parser"""
        logger.info("Document parser service initialized")
        if OCR_AVAILABLE:
            logger.info("✓ OCR support available (pytesseract)")
        else:
            logger.warning("⚠ OCR not available - install tesseract for scanned document support")

    # =========================================================================
    # Main Parsing Methods
    # =========================================================================

    def parse_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse document and extract text

        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type (optional)

        Returns:
            Dictionary with extracted text and metadata
        """
        # Determine file type
        file_ext = Path(filename).suffix.lower()

        logger.info(f"Parsing document: {filename} (type: {file_ext})")

        try:
            if file_ext == '.pdf' or content_type == 'application/pdf':
                return self.parse_pdf(file_content, filename)

            elif file_ext in ['.docx', '.doc'] or content_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]:
                return self.parse_docx(file_content, filename)

            elif file_ext == '.txt' or content_type == 'text/plain':
                return self.parse_text(file_content, filename)

            else:
                # Try parsing as text by default
                logger.warning(f"Unknown file type {file_ext}, attempting text parsing")
                return self.parse_text(file_content, filename)

        except Exception as e:
            logger.error(f"Document parsing failed for {filename}: {e}")
            raise

    # =========================================================================
    # PDF Parsing
    # =========================================================================

    def parse_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse PDF document and extract text

        Args:
            file_content: PDF file bytes
            filename: Filename

        Returns:
            Parsed document with text and metadata
        """
        try:
            # Open PDF from bytes
            pdf_stream = io.BytesIO(file_content)
            pdf_document = fitz.open(stream=pdf_stream, filetype="pdf")

            # Extract text from all pages
            full_text = ""
            pages_data = []

            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]

                # Extract text
                page_text = page.get_text()

                # Page metadata
                page_info = {
                    "page_number": page_num + 1,
                    "text": page_text,
                    "word_count": len(page_text.split()),
                    "char_count": len(page_text)
                }

                pages_data.append(page_info)
                full_text += page_text + "\n\n"

            # Document metadata
            metadata = pdf_document.metadata or {}

            pdf_document.close()

            logger.info(f"✓ PDF parsed: {filename} ({len(pdf_document)} pages, {len(full_text)} chars)")

            return {
                "text": full_text.strip(),
                "filename": filename,
                "format": "pdf",
                "page_count": len(pages_data),
                "pages": pages_data,
                "metadata": {
                    "title": metadata.get("title"),
                    "author": metadata.get("author"),
                    "subject": metadata.get("subject"),
                    "creator": metadata.get("creator"),
                    "producer": metadata.get("producer"),
                },
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }

        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")

            # Try OCR if text extraction failed
            if OCR_AVAILABLE:
                logger.info("Attempting OCR extraction...")
                return self.parse_pdf_with_ocr(file_content, filename)
            else:
                raise

    def parse_pdf_with_ocr(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse scanned PDF using OCR

        Args:
            file_content: PDF file bytes
            filename: Filename

        Returns:
            OCR-extracted text and metadata
        """
        if not OCR_AVAILABLE:
            raise RuntimeError("OCR not available - install pytesseract and tesseract")

        try:
            pdf_stream = io.BytesIO(file_content)
            pdf_document = fitz.open(stream=pdf_stream, filetype="pdf")

            full_text = ""
            pages_data = []

            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]

                # Render page to image
                pix = page.get_pixmap(dpi=300)  # High DPI for better OCR
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))

                # Perform OCR
                page_text = pytesseract.image_to_string(image, lang='eng')

                pages_data.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "ocr": True
                })

                full_text += page_text + "\n\n"

            pdf_document.close()

            logger.info(f"✓ PDF OCR complete: {filename} ({len(pages_data)} pages)")

            return {
                "text": full_text.strip(),
                "filename": filename,
                "format": "pdf_ocr",
                "page_count": len(pages_data),
                "pages": pages_data,
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }

        except Exception as e:
            logger.error(f"PDF OCR failed: {e}")
            raise

    # =========================================================================
    # DOCX Parsing
    # =========================================================================

    def parse_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse Microsoft Word DOCX document

        Args:
            file_content: DOCX file bytes
            filename: Filename

        Returns:
            Parsed document with text and structure
        """
        try:
            # Open DOCX from bytes
            docx_stream = io.BytesIO(file_content)
            document = DocxDocument(docx_stream)

            # Extract text from paragraphs
            paragraphs = []
            full_text = ""

            for para in document.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })
                    full_text += para.text + "\n"

            # Extract text from tables
            tables_data = []
            for table in document.tables:
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_text += row_text + "\n"

                tables_data.append(table_text)
                full_text += table_text + "\n"

            # Document properties
            props = document.core_properties

            logger.info(f"✓ DOCX parsed: {filename} ({len(paragraphs)} paragraphs, {len(tables_data)} tables)")

            return {
                "text": full_text.strip(),
                "filename": filename,
                "format": "docx",
                "paragraph_count": len(paragraphs),
                "paragraphs": paragraphs,
                "table_count": len(tables_data),
                "tables": tables_data,
                "metadata": {
                    "title": props.title,
                    "author": props.author,
                    "subject": props.subject,
                    "created": props.created.isoformat() if props.created else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                },
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }

        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise

    # =========================================================================
    # Text Parsing
    # =========================================================================

    def parse_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse plain text file

        Args:
            file_content: Text file bytes
            filename: Filename

        Returns:
            Parsed text
        """
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                text = file_content.decode('latin-1', errors='ignore')

            logger.info(f"✓ Text parsed: {filename} ({len(text)} chars)")

            return {
                "text": text.strip(),
                "filename": filename,
                "format": "text",
                "word_count": len(text.split()),
                "char_count": len(text),
                "line_count": len(text.splitlines())
            }

        except Exception as e:
            logger.error(f"Text parsing failed: {e}")
            raise

    # =========================================================================
    # Section Detection
    # =========================================================================

    def detect_sections(self, text: str) -> Dict[str, str]:
        """
        Detect common clinical note sections

        Args:
            text: Document text

        Returns:
            Dictionary mapping section names to text
        """
        sections = {}

        # Common section headers
        section_patterns = {
            "chief_complaint": ["CHIEF COMPLAINT", "CC:"],
            "history": ["HISTORY OF PRESENT ILLNESS", "HPI:", "HISTORY:"],
            "past_medical_history": ["PAST MEDICAL HISTORY", "PMH:"],
            "medications": ["MEDICATIONS:", "CURRENT MEDICATIONS", "MEDS:"],
            "allergies": ["ALLERGIES:", "DRUG ALLERGIES"],
            "physical_exam": ["PHYSICAL EXAMINATION", "PHYSICAL EXAM", "EXAM:"],
            "assessment": ["ASSESSMENT:", "IMPRESSION:"],
            "plan": ["PLAN:", "TREATMENT PLAN"],
            "labs": ["LABORATORY", "LABS:", "LAB RESULTS"],
            "imaging": ["IMAGING:", "RADIOLOGY", "MRI", "CT SCAN"],
        }

        # Simple section detection (can be enhanced with ML)
        lines = text.split('\n')
        current_section = None

        for line in lines:
            line_upper = line.strip().upper()

            # Check if line is a section header
            for section_name, patterns in section_patterns.items():
                if any(pattern in line_upper for pattern in patterns):
                    current_section = section_name
                    sections[section_name] = ""
                    break

            # Add line to current section
            if current_section and line.strip():
                sections[current_section] += line + "\n"

        logger.info(f"Detected {len(sections)} sections in document")
        return sections


# =============================================================================
# Global Parser Instance
# =============================================================================

_document_parser: Optional[DocumentParserService] = None


def get_document_parser() -> DocumentParserService:
    """Get or create document parser instance"""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParserService()
    return _document_parser


# =============================================================================
# Public API Functions
# =============================================================================

def parse_file(
    file_content: bytes,
    filename: str,
    content_type: Optional[str] = None
) -> Dict[str, Any]:
    """
    Parse document file and extract text

    Args:
        file_content: File bytes
        filename: Original filename
        content_type: MIME type

    Returns:
        Parsed document with text and metadata
    """
    parser = get_document_parser()
    return parser.parse_document(file_content, filename, content_type)


def detect_document_sections(text: str) -> Dict[str, str]:
    """Detect clinical note sections"""
    parser = get_document_parser()
    return parser.detect_sections(text)
